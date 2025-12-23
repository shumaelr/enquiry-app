"""
Extractors module for reading PDF and Word documents.
"""
import pypdf
from docx import Document
import os


def extract_pdf_text(filepath: str) -> str | None:
    """
    Extracts text content from a PDF file.
    
    Args:
        filepath: Path to the PDF file.
        
    Returns:
        Extracted text or None if an error occurs.
    """
    try:
        if not os.path.exists(filepath):
            print(f"PDF file not found: {filepath}")
            return None
            
        reader = pypdf.PdfReader(filepath)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        if not text.strip():
            print(f"Warning: No text extracted from PDF {filepath}")
            return None
            
        return text
        
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
        return None


def extract_docx_text(filepath: str) -> str | None:
    """
    Extracts text content from a Word document (.docx).
    
    Args:
        filepath: Path to the Word document.
        
    Returns:
        Extracted text or None if an error occurs.
    """
    try:
        if not os.path.exists(filepath):
            print(f"DOCX file not found: {filepath}")
            return None
            
        doc = Document(filepath)
        text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        result = "\n".join(text)
        
        if not result.strip():
            print(f"Warning: No text extracted from DOCX {filepath}")
            return None
            
        return result
        
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
        return None
